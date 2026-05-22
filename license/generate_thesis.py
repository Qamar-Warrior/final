"""
Generate Bachelor Thesis .docx for the Uzbekistan License Plate Recognition project.
Follows the exact formatting rules extracted from Bachelor Thesis Format.docx.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import copy

doc = Document()

# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)

def set_single_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

def add_run(paragraph, text, bold=False, italic=False, size=14, font_name="Times New Roman"):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    return run

def body_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              indent_first=1.5, ls=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, ls)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    add_run(p, text, bold=bold, italic=italic)
    return p

def chapter_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, text, bold=True)
    return p

def section_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, text, bold=True)
    return p

def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, text, italic=True)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
    return p

def page_break():
    doc.add_page_break()

# ─────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────

def title_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14):
    p = doc.add_paragraph()
    p.alignment = align
    set_single_spacing(p)
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, text, bold=bold, size=size)
    return p

title_para("MILLAT UMIDI UNIVERSITY", bold=True)
title_para("")
title_para("SURNAME   NAME")
title_para("")
title_para("UZBEKISTAN LICENSE PLATE RECOGNITION SYSTEM", bold=True)
title_para("")
title_para('60610600 – Software Engineering\n("Dasturiy injiniringi")')
title_para("")
title_para("Bachelor Degree Thesis", size=24)
title_para("")
title_para("")

sup = doc.add_paragraph()
sup.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_single_spacing(sup)
sup.paragraph_format.first_line_indent = Cm(0)
add_run(sup, "Supervisor: ____________,", bold=True)

sup2 = doc.add_paragraph()
sup2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_single_spacing(sup2)
sup2.paragraph_format.first_line_indent = Cm(0)
add_run(sup2, "IT Department, Senior teacher")

rev = doc.add_paragraph()
rev.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_single_spacing(rev)
rev.paragraph_format.first_line_indent = Cm(0)
add_run(rev, "Reviewer: ______________,", bold=True)

rev2 = doc.add_paragraph()
rev2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_single_spacing(rev2)
rev2.paragraph_format.first_line_indent = Cm(0)
add_run(rev2, "IT Department, Senior teacher")

doc.add_paragraph()

adm = doc.add_paragraph()
adm.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_single_spacing(adm)
adm.paragraph_format.first_line_indent = Cm(0)
add_run(adm, '"Admitted to defense"\nDean of the Information Technologies Faculty\n_________M.M. Pirnazarov\n"___" ___________ 2025')

doc.add_paragraph()
title_para("Tashkent – 2025")

page_break()

# ─────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────

chapter_heading("TABLE OF CONTENTS")

toc_entries = [
    ("Introduction", "3", False),
    ("CHAPTER 1. LICENSE PLATE RECOGNITION — OVERVIEW AND BACKGROUND", "", True),
    ("  1.1 History and importance of license plate recognition", "5", False),
    ("  1.2 Types of license plate recognition systems", "8", False),
    ("  1.3 Uzbekistan license plate standard", "11", False),
    ("CHAPTER 2. TECHNOLOGIES AND TOOLS", "", True),
    ("  2.1 YOLOv8 object detection framework", "14", False),
    ("  2.2 EasyOCR optical character recognition", "18", False),
    ("  2.3 FastAPI and SQLite", "22", False),
    ("  2.4 OpenCV image processing", "26", False),
    ("CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION", "", True),
    ("  3.1 System architecture and pipeline", "30", False),
    ("  3.2 Detection module", "34", False),
    ("  3.3 Recognition and validation modules", "38", False),
    ("  3.4 REST API and database", "43", False),
    ("CONCLUSION", "48", True),
    ("LIST OF USED LITERATURE", "50", True),
    ("SOURCE CODE STRUCTURE AND DESCRIPTION", "51", True),
]

for label, page, bold in toc_entries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    if page:
        add_run(p, f"{label} {'.' * max(1, 70 - len(label) - len(page))} {page}", bold=bold)
    else:
        add_run(p, label, bold=bold)

page_break()

# ─────────────────────────────────────────────
# INTRODUCTION
# ─────────────────────────────────────────────

chapter_heading("INTRODUCTION")

body_para(
    "In the modern era of intelligent transportation systems, the automated recognition of "
    "vehicle license plates has become an indispensable tool for traffic management, security "
    "surveillance, parking automation, and law enforcement. License Plate Recognition (LPR), "
    "also referred to as Automatic Number Plate Recognition (ANPR), employs computer vision "
    "and machine learning techniques to detect, extract, and interpret the alphanumeric "
    "characters displayed on vehicle registration plates."
)

body_para(
    "The Republic of Uzbekistan has experienced rapid growth in vehicle ownership over the "
    "past decade, placing greater demand on traffic authorities to monitor and manage road "
    "usage effectively. Uzbekistan uses a nationally standardized plate format — two-digit "
    "region code, one letter, three digits, and two letters (e.g., 01A123BC) — which provides "
    "a well-defined structure that can be exploited by automated recognition systems."
)

body_para(
    "This thesis presents the design and implementation of an end-to-end Uzbekistan License "
    "Plate Recognition System. The system integrates a fine-tuned YOLOv8 deep-learning model "
    "for plate localization, EasyOCR for character extraction, and a rule-based validator "
    "that applies both format and region-code constraints to ensure only legitimate plates "
    "are recorded. A FastAPI-based REST interface exposes the pipeline to external clients, "
    "and an SQLite database with Write-Ahead Logging (WAL) stores detection records for "
    "later querying and analysis."
)

body_para(
    "The work is organized as follows. Chapter 1 surveys the background of license plate "
    "recognition, including historical developments, system types, and the specifics of the "
    "Uzbekistan plate standard. Chapter 2 describes the key technologies employed. Chapter 3 "
    "details the system architecture, module design, and implementation decisions. The "
    "Conclusion summarizes findings and suggests directions for future improvement."
)

page_break()

# ─────────────────────────────────────────────
# CHAPTER 1
# ─────────────────────────────────────────────

chapter_heading("CHAPTER 1. LICENSE PLATE RECOGNITION — OVERVIEW AND BACKGROUND")

section_heading("1.1 History and Importance of License Plate Recognition")

body_para(
    "License plate recognition has its origins in the early 1970s when researchers at the "
    "Police Scientific Development Branch in the United Kingdom began exploring optical "
    "character recognition applied to vehicle plates for law-enforcement purposes. Early "
    "systems relied on specialized infrared cameras, fixed thresholding, and hand-crafted "
    "feature extractors, and could only operate under tightly controlled lighting conditions."
)

body_para(
    "The 1990s brought a surge of interest driven by the availability of affordable frame "
    "grabbers and faster general-purpose processors. Rule-based segmentation approaches — "
    "edge detection, connected-component analysis, and template matching — became the "
    "dominant paradigm. Systems of this era achieved useful accuracy on high-quality images "
    "but degraded significantly with blur, occlusion, or non-standard plate fonts."
)

body_para(
    "The deep-learning revolution of the 2010s transformed the field. Convolutional neural "
    "networks (CNNs) demonstrated that a single model trained end-to-end on large annotated "
    "datasets could learn both localization and character recognition far more robustly than "
    "any handcrafted pipeline. Today, LPR systems are deployed at toll booths, border "
    "crossings, shopping-mall car parks, and smart-city traffic nodes worldwide."
)

body_para(
    "The importance of LPR extends beyond convenience. Traffic authorities use it to enforce "
    "speed limits, detect stolen vehicles, and manage congestion pricing. Security agencies "
    "monitor border crossings and sensitive zones. Parking operators automate entry and exit "
    "without human intervention. Each of these applications demands high accuracy, low "
    "latency, and reliable operation across diverse environmental conditions."
)

section_heading("1.2 Types of License Plate Recognition Systems")

body_para(
    "LPR systems can be classified along several dimensions: deployment environment, "
    "processing architecture, and the detection methodology used.",
    bold=True
)

body_para(
    "Fixed versus mobile systems. Fixed installations mount cameras at known, controlled "
    "positions — gantries above motorways, entrance barriers, or roadside poles. The "
    "controlled viewpoint simplifies detection geometry. Mobile systems, fitted to police "
    "vehicles or drones, must handle arbitrary viewing angles, variable distances, and "
    "rapidly changing backgrounds."
)

body_para(
    "Edge versus cloud processing. With the proliferation of embedded GPUs and neural "
    "processing units, many modern LPR deployments perform inference on the device itself, "
    "reducing latency and bandwidth requirements. Cloud-based architectures offload "
    "computation to remote servers, enabling more powerful models but introducing network "
    "dependency and privacy considerations."
)

body_para(
    "Traditional versus deep-learning pipelines. Traditional pipelines decompose the problem "
    "into independent stages — plate detection (edge/blob detection), character segmentation "
    "(projection profiles), and character recognition (k-NN or SVM on hand-crafted features). "
    "Deep-learning pipelines, by contrast, replace one or more stages with neural network "
    "components that learn representations directly from data."
)

figure_caption("Figure 1.2.1. Taxonomy of license plate recognition system architectures.")

body_para(
    "The present system follows a hybrid deep-learning approach: YOLOv8 — a real-time object "
    "detector — handles plate localization, while EasyOCR — a transformer-based OCR engine — "
    "reads the extracted plate crop. This combination balances accuracy, inference speed, and "
    "ease of deployment."
)

section_heading("1.3 Uzbekistan License Plate Standard")

body_para(
    "Uzbekistan introduced its current national registration plate standard in 1994, "
    "following independence from the Soviet Union. The standard defines both the physical "
    "plate dimensions and the encoding scheme used to assign registration numbers."
)

body_para(
    "Physical specifications. Standard passenger-vehicle plates measure 520 mm × 112 mm "
    "(approximately 4.6:1 aspect ratio). Plates carry black characters on a white background "
    "for most civilian categories, with distinct color schemes for diplomatic, government, "
    "and temporary registrations."
)

body_para(
    "Encoding format. Each plate encodes exactly eight characters following the pattern "
    "DD L DDD LL, where D denotes a decimal digit and L denotes an uppercase Latin letter. "
    "The first two digits identify the region of registration, the single letter at position "
    "three indicates the series within that region, the three-digit block encodes a sequential "
    "number, and the final two letters complete the identifier."
)

figure_caption("Figure 1.3.1. Structure of an Uzbekistan vehicle registration plate.")

body_para(
    "Region codes. Fourteen official region codes are currently in use, ranging from 01 "
    "(Tashkent city) to 80 (Republic of Karakalpakstan). Each code corresponds to one of "
    "the twelve provinces (viloyats), the capital city, or the autonomous republic. The "
    "system developed in this thesis encodes all fourteen valid region codes and uses them "
    "as an additional validation layer: a plate whose two-digit prefix is not one of the "
    "fourteen codes is flagged as invalid regardless of whether the remaining characters "
    "match the format pattern."
)

body_para(
    "OCR challenges. The Uzbekistan plate font uses characters that are susceptible to "
    "common OCR confusions — the letter O and the digit 0, the letter I and the digit 1, "
    "the letter S and the digit 5, and several others. The validator incorporates a "
    "correction table that attempts to remap misrecognized characters to their most likely "
    "intended counterpart before applying the format check."
)

page_break()

# ─────────────────────────────────────────────
# CHAPTER 2
# ─────────────────────────────────────────────

chapter_heading("CHAPTER 2. TECHNOLOGIES AND TOOLS")

section_heading("2.1 YOLOv8 Object Detection Framework")

body_para(
    "YOLO (You Only Look Once) is a family of real-time object detection models first "
    "introduced by Redmon et al. in 2016. Unlike two-stage detectors (e.g., Faster R-CNN) "
    "that generate region proposals before classifying them, YOLO formulates object "
    "detection as a single regression problem: a single forward pass through the network "
    "simultaneously predicts bounding boxes and class probabilities for all objects in an "
    "image."
)

body_para(
    "YOLOv8, released by Ultralytics in January 2023, is the eighth major iteration of the "
    "architecture. It introduces anchor-free detection heads, a decoupled head design that "
    "separates classification and regression branches, and a new backbone (C2f modules) "
    "derived from YOLOv7's E-ELAN. These changes improve accuracy at all scales while "
    "maintaining inference speeds suitable for real-time video processing."
)

figure_caption("Figure 2.1.1. YOLOv8 architecture overview — backbone, neck, and detection head.")

body_para(
    "The Ultralytics Python library wraps the model in a convenient API. A single call to "
    "YOLO('model.pt') loads the weights, and model(frame, conf=0.45) runs inference and "
    "returns structured result objects containing bounding-box coordinates, confidence "
    "scores, and class labels. The library handles batching, GPU acceleration via CUDA, "
    "and model export to ONNX, TensorRT, and CoreML formats."
)

body_para(
    "In the present system, a custom-trained YOLOv8 model specialized for Uzbekistan license "
    "plates is loaded at startup. The model was fine-tuned on a dataset of annotated plate "
    "images captured under diverse lighting and weather conditions. A fallback to the "
    "general-purpose YOLOv8n (COCO) model is provided for environments where the custom "
    "weights are unavailable; in fallback mode, aspect-ratio filtering (2.0 to 8.0 width-to-"
    "height ratio) restricts detections to plate-shaped regions."
)

section_heading("2.2 EasyOCR Optical Character Recognition")

body_para(
    "EasyOCR is an open-source Python library built by JaidedAI that provides ready-to-use "
    "OCR for over 80 languages. Internally it uses a CRAFT (Character Region Awareness for "
    "Text Detection) network for text-region detection and a CRNN (Convolutional Recurrent "
    "Neural Network) with a CTC (Connectionist Temporal Classification) decoder for "
    "recognition. This architecture handles variable-length text without requiring character-"
    "level segmentation."
)

body_para(
    "Initialization requires specifying the target language list and whether to use GPU "
    "acceleration. For Uzbekistan plates the English ('en') language pack is sufficient "
    "because all plate characters are Latin digits and uppercase letters. The reader object "
    "is created once at application startup to amortize the high initialization cost (model "
    "loading and JIT compilation) across all subsequent inference calls."
)

figure_caption("Figure 2.2.1. EasyOCR CRNN recognition pipeline with CTC decoder.")

body_para(
    "An allowlist parameter restricts recognized characters to the set "
    "ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789, eliminating spurious detections of punctuation "
    "or symbols that cannot appear on a legitimate plate. Before passing the cropped plate "
    "image to EasyOCR, the system applies a four-step preprocessing chain: grayscale "
    "conversion, bicubic resizing to a fixed height of 128 pixels, Otsu binarization to "
    "normalize contrast, and a 3×3 unsharp-masking kernel to sharpen character edges."
)

section_heading("2.3 FastAPI and SQLite")

body_para(
    "FastAPI is a modern, high-performance Python web framework for building APIs with "
    "automatic OpenAPI documentation generation. It is built on top of Starlette (for the "
    "ASGI core) and Pydantic (for request/response validation). FastAPI's asynchronous "
    "request handling allows it to serve many concurrent connections efficiently without "
    "multi-threading complexity."
)

body_para(
    "Pydantic models define the schema for every API request and response. The framework "
    "automatically validates incoming data, rejects malformed requests with structured error "
    "messages, and serializes outgoing responses to JSON. Interactive documentation at "
    "/docs (Swagger UI) and /redoc is generated automatically from the type annotations, "
    "enabling rapid API exploration during development and integration."
)

figure_caption("Figure 2.3.1. FastAPI automatic Swagger UI documentation interface.")

body_para(
    "SQLite is a self-contained, serverless relational database engine embedded in the "
    "application process. It stores all data in a single file (plates.db), making deployment "
    "and backup trivial. The system enables WAL (Write-Ahead Logging) mode, which allows "
    "multiple concurrent readers and a single writer to operate simultaneously without "
    "blocking — essential for the REST API (reads) and background video processing (writes) "
    "to coexist."
)

body_para(
    "Two tables are defined: detections stores every detection attempt including invalid "
    "plates, enabling analysis of OCR failure modes. valid_plates stores only confirmed "
    "valid plate reads for fast lookup by enforcement systems. Indexes on plate_text, "
    "detected_at, and is_valid support the most common query patterns."
)

section_heading("2.4 OpenCV Image Processing")

body_para(
    "OpenCV (Open Source Computer Vision Library) is a cross-platform library of over 2,500 "
    "optimized algorithms for image and video analysis. The Python bindings (opencv-python) "
    "expose the full C++ API with NumPy array integration, making it the natural choice "
    "for frame acquisition and preprocessing in this project."
)

body_para(
    "In the detection module, cv2.imread() and cv2.VideoCapture() handle file-based and "
    "streaming input respectively. The VideoCapture object exposes individual frames as "
    "NumPy arrays in BGR color order. Frame skipping — processing every Nth frame rather "
    "than every frame — is implemented by checking the modulo of the frame index against "
    "the configurable VIDEO_FRAME_SKIP parameter (default 3), reducing CPU load by "
    "approximately 67 percent while preserving detection coverage."
)

figure_caption("Figure 2.4.1. OpenCV bounding-box annotation overlaid on a detection result.")

body_para(
    "The draw_boxes() method of the detector class uses cv2.rectangle() and cv2.putText() "
    "to render green bounding boxes and plate-text labels on annotated output frames. This "
    "visualization path is available in the CLI's --show mode for development and "
    "demonstration purposes."
)

page_break()

# ─────────────────────────────────────────────
# CHAPTER 3
# ─────────────────────────────────────────────

chapter_heading("CHAPTER 3. SYSTEM DESIGN AND IMPLEMENTATION")

section_heading("3.1 System Architecture and Pipeline")

body_para(
    "The system is organized into four loosely coupled layers: the core processing pipeline, "
    "the database layer, the REST API layer, and the CLI entry point. Each layer communicates "
    "through well-defined interfaces, allowing components to be tested and replaced "
    "independently."
)

figure_caption("Figure 3.1.1. High-level system architecture diagram.")

body_para(
    "The core pipeline (core/pipeline.py) is the central orchestrator. On initialization it "
    "instantiates a PlateDetector and a TextRecognizer, loading model weights into memory "
    "once. Two public methods — process_image() and process_video() — expose the pipeline "
    "to callers. Both methods ultimately delegate to the private _process_frame() method, "
    "which executes the three-stage sequence: detect → recognize → validate."
)

body_para(
    "Configuration is centralized in config.py, which reads all tuneable parameters from "
    "environment variables with sensible defaults. This design allows operators to adjust "
    "confidence thresholds, frame-skip rates, GPU usage, and database paths without "
    "modifying source code — a requirement for container-based deployments."
)

body_para(
    "The data flow for a single frame is as follows: the raw BGR frame is passed to "
    "PlateDetector.detect(), which returns a list of detection dictionaries each containing "
    "the bounding box, confidence score, and cropped plate image. Each crop is passed to "
    "TextRecognizer.recognize(), which returns the concatenated raw text and a mean "
    "confidence. The raw text is passed to validate_plate(), which returns a boolean "
    "validity flag and the normalized plate string. A combined confidence is computed as "
    "the geometric mean of YOLO and OCR scores to balance the two sources of uncertainty."
)

section_heading("3.2 Detection Module")

body_para(
    "The PlateDetector class (core/detector.py) wraps the Ultralytics YOLO API. At "
    "construction time it attempts to load a custom plate-specific model from the path "
    "specified in config.MODEL_PATH. If the file does not exist, it falls back to the "
    "general-purpose YOLOv8n model and activates aspect-ratio filtering."
)

body_para(
    "The detect() method runs model inference on a full BGR frame and iterates over the "
    "result objects returned by the YOLO API. For each detected bounding box it extracts "
    "the pixel coordinates using box.xyxy[0], clips the crop from the frame using NumPy "
    "array slicing, and appends the detection dictionary to the result list. Empty crops "
    "(zero-area bounding boxes) are skipped to prevent downstream errors."
)

figure_caption("Figure 3.2.1. Sample detection output — bounding box overlaid on vehicle image.")

body_para(
    "In fallback mode two additional filters are applied. The width-to-height ratio must "
    "fall between 2.0 and 8.0, consistent with the physical dimensions of vehicle plates. "
    "The bounding box area must not exceed 15 percent of the total frame area, eliminating "
    "detections of large objects (vehicle bodies, road signs) that would otherwise trigger "
    "false positives on the general-purpose model."
)

section_heading("3.3 Recognition and Validation Modules")

body_para(
    "The TextRecognizer class (core/recognizer.py) wraps the EasyOCR Reader. The "
    "constructor initializes the reader with the configured language list and performs a "
    "warm-up pass on a blank white image. This warm-up triggers JIT compilation inside "
    "EasyOCR, ensuring the first real inference call does not incur an unexpected latency "
    "spike."
)

body_para(
    "The recognize() method applies the four-step preprocessing pipeline described in "
    "Section 2.2 before calling reader.readtext(). All recognized text segments are "
    "concatenated without separators — consistent with the Uzbekistan plate format, which "
    "has no spaces or delimiters — and the per-segment confidences are averaged. The method "
    "returns a (raw_text, mean_confidence) tuple."
)

figure_caption("Figure 3.3.1. Preprocessing stages applied to a plate crop before OCR.")

body_para(
    "The validator module (core/validator.py) applies a two-level check. First, a compiled "
    "regular expression UZ_PLATE_RE = re.compile(r'^[0-9]{2}[A-Z][0-9]{3}[A-Z]{2}$') "
    "verifies the structural format. Second, the two-digit prefix is looked up in the "
    "VALID_REGIONS set of fourteen official codes. If both checks pass, the plate is "
    "accepted. If the structural check fails, the module attempts OCR correction: for each "
    "character position that should be a digit but contains a letter (or vice versa), it "
    "consults a lookup table of common OCR misreads (e.g., 'O' → '0', 'I' → '1', "
    "'S' → '5') and substitutes the corrected character before re-evaluating."
)

body_para(
    "Video deduplication prevents the same plate from being written to the database "
    "repeatedly as it remains in the camera's field of view across multiple frames. A "
    "dictionary maps each recognized plate_text to the datetime of its most recent "
    "detection. A new result is considered a duplicate if the same plate was last seen "
    "within the configurable DEDUP_WINDOW_SECONDS interval (default 2 seconds). The cache "
    "is pruned when it exceeds 1,000 entries to prevent unbounded memory growth."
)

section_heading("3.4 REST API and Database")

body_para(
    "The API layer (api/app.py, api/routes.py) exposes five endpoint groups. The health "
    "endpoint GET /health returns the current operational status and database path. Two "
    "detection endpoints accept file uploads: POST /detect/image processes a single image "
    "synchronously and returns a JSON array of detections; POST /detect/video processes an "
    "uploaded video file and streams results as newline-delimited JSON (NDJSON) so the "
    "client can consume results incrementally without waiting for the full file to be "
    "processed."
)

figure_caption("Figure 3.4.1. REST API endpoint summary and data flow.")

body_para(
    "Three query endpoints support retrieval of stored records: GET /detections lists "
    "recent detections with pagination and an optional valid_only filter; "
    "GET /detections/search?q=01A123BC performs a case-insensitive partial match on the "
    "plate_text column; GET /detections/{id} retrieves a single record by primary key. "
    "GET /stats returns aggregate counts of total, valid, and invalid detections grouped "
    "by source."
)

body_para(
    "The DatabaseManager class (db/manager.py) provides a thin wrapper around the sqlite3 "
    "standard-library module. All queries use parameterized statements to prevent SQL "
    "injection. The connection is opened with check_same_thread=False to allow the FastAPI "
    "async runtime — which may dispatch handlers across OS threads — to share the "
    "connection safely. WAL mode is enabled at startup via PRAGMA journal_mode=WAL, "
    "allowing concurrent reads during write operations."
)

figure_caption("Figure 3.4.2. Entity-relationship diagram of the database schema.")

body_para(
    "The CLI entry point (main.py) uses Python's argparse module to expose four "
    "subcommands: image, video, serve, and query. The image and video subcommands "
    "instantiate the pipeline and database, process the specified file, and print JSON "
    "results to standard output — enabling easy integration with shell scripts and "
    "downstream tools. The serve subcommand starts the Uvicorn ASGI server. The query "
    "subcommand allows direct database interrogation from the command line without "
    "starting the API server."
)

page_break()

# ─────────────────────────────────────────────
# CONCLUSION
# ─────────────────────────────────────────────

chapter_heading("CONCLUSION")

body_para(
    "This thesis has presented the full design and implementation of an end-to-end "
    "Uzbekistan License Plate Recognition System. The system integrates three well-established "
    "open-source technologies — YOLOv8 for plate localization, EasyOCR for character "
    "recognition, and FastAPI with SQLite for API serving and persistent storage — into a "
    "cohesive pipeline that can process both static images and real-time video streams."
)

body_para(
    "The modular architecture cleanly separates concerns: the detector, recognizer, validator, "
    "database manager, and API router are each encapsulated in a single class with a minimal "
    "public interface. This separation facilitates independent testing, replacement of "
    "components with improved implementations, and deployment in diverse environments "
    "ranging from a developer laptop to a containerized cloud service."
)

body_para(
    "The two-level validation strategy — structural regex plus regional code lookup plus "
    "OCR-error correction — significantly improves the precision of stored records compared "
    "to raw OCR output alone. The deduplication mechanism prevents database bloat during "
    "video processing while still capturing every distinct plate that appears in a stream."
)

body_para(
    "Several directions for future work are identified. First, the custom YOLOv8 model "
    "could be retrained on a larger and more diverse dataset covering night-time, rain, "
    "and heavy-motion scenarios to improve recall in adverse conditions. Second, real-time "
    "streaming via WebSocket or Server-Sent Events would allow the API to push detections "
    "to connected clients without polling. Third, integration with an external vehicle "
    "registration database would enable the system to flag unregistered or flagged plates "
    "automatically. Fourth, a web-based dashboard providing live map visualization of "
    "detection events would increase operational utility for traffic-management centers."
)

body_para(
    "In summary, the system demonstrates that a compact, dependency-light Python application "
    "can deliver production-quality license plate recognition for the Uzbekistan context, "
    "providing a solid foundation for further research and operational deployment."
)

page_break()

# ─────────────────────────────────────────────
# LIST OF USED LITERATURE
# ─────────────────────────────────────────────

chapter_heading("LIST OF USED LITERATURE")

refs = [
    "Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You Only Look Once: Unified, Real-Time Object Detection. CVPR 2016.",
    "Jocher, G., Chaurasia, A., & Qiu, J. (2023). Ultralytics YOLOv8. GitHub Repository. https://github.com/ultralytics/ultralytics",
    "JaidedAI. (2020). EasyOCR: Ready-to-use OCR with 80+ supported languages. GitHub Repository. https://github.com/JaidedAI/EasyOCR",
    "Baek, Y., Lee, B., Han, D., Yun, S., & Lee, H. (2019). Character Region Awareness for Text Detection. CVPR 2019.",
    "Shi, B., Bai, X., & Yao, C. (2016). An End-to-End Trainable Neural Network for Image-based Sequence Recognition. IEEE TPAMI.",
    "Tiangolo, S. (2018). FastAPI: The Modern, Fast Web Framework for Python. https://fastapi.tiangolo.com",
    "Bradski, G. (2000). The OpenCV Library. Dr. Dobb's Journal of Software Tools.",
    "Hipp, R. D. (2000). SQLite: A Self-Contained, Serverless, Zero-Configuration, Transactional SQL Database Engine. https://www.sqlite.org",
    "Patel, C., Shah, D., & Patel, A. (2013). Automatic Number Plate Recognition System (ANPR): A Survey. International Journal of Computer Applications.",
    "Jain, A. K., & Yu, B. (1998). Automatic Text Location in Images and Video Frames. Pattern Recognition, 31(12).",
    "Özpoyraz, M., & Tosun, Ö. (2020). Deep Learning Based License Plate Recognition: A Survey. Engineering Science and Technology.",
    "Wikipedia contributors. (2024). Automatic number-plate recognition. Wikipedia, The Free Encyclopedia.",
    "Ultralytics Documentation. (2024). YOLOv8 Model Training and Inference. https://docs.ultralytics.com",
    "Python Software Foundation. (2024). sqlite3 — DB-API 2.0 interface for SQLite databases. https://docs.python.org/3/library/sqlite3.html",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    add_run(p, f"{i}. {ref}")

page_break()

# ─────────────────────────────────────────────
# SOURCE CODE STRUCTURE AND DESCRIPTION
# ─────────────────────────────────────────────

chapter_heading("SOURCE CODE STRUCTURE AND DESCRIPTION")

body_para("Project directory structure:", bold=True, indent_first=0)

code_block(
    "license/\n"
    "├── main.py                  # CLI entry point (image / video / serve / query)\n"
    "├── config.py                # Environment-variable-driven configuration\n"
    "├── requirements.txt         # Python dependency list\n"
    "├── plates.db                # SQLite database file\n"
    "├── core/\n"
    "│   ├── detector.py          # YOLOv8-based plate localization\n"
    "│   ├── recognizer.py        # EasyOCR-based character extraction\n"
    "│   ├── validator.py         # Format + region-code validation\n"
    "│   └── pipeline.py          # Orchestrator: detect → recognize → validate\n"
    "├── api/\n"
    "│   ├── app.py               # FastAPI application factory\n"
    "│   ├── routes.py            # API endpoint definitions\n"
    "│   └── schemas.py           # Pydantic request/response models\n"
    "├── db/\n"
    "│   ├── manager.py           # SQLite DatabaseManager class\n"
    "│   └── schema.sql           # CREATE TABLE statements\n"
    "├── models/\n"
    "│   └── plate_detector.pt    # Fine-tuned YOLOv8 model weights\n"
    "└── tests/\n"
    "    ├── test_validator.py\n"
    "    ├── test_db.py\n"
    "    └── test_pipeline_synthetic.py"
)

body_para("Core pipeline entry point (main.py — process_image command):", bold=True, indent_first=0)

code_block(
    "def cmd_image(args):\n"
    "    pipeline = _make_pipeline()\n"
    "    db = _make_db() if not args.no_save else None\n"
    "    results = pipeline.process_image(args.path)\n"
    "    for result in results:\n"
    "        should_save = db and result.confidence >= config.MIN_SAVE_CONFIDENCE\n"
    "        record_id = db.save_detection(result) if should_save else -1\n"
    "        if should_save and result.is_valid:\n"
    "            db.save_valid_plate(result)\n"
    "        print(json.dumps(_result_to_dict(result, record_id)))"
)

body_para("Plate validation logic (core/validator.py):", bold=True, indent_first=0)

code_block(
    "UZ_PLATE_RE = re.compile(r'^[0-9]{2}[A-Z][0-9]{3}[A-Z]{2}$')\n\n"
    "VALID_REGIONS = {\n"
    "    '01', '10', '20', '25', '30', '35',\n"
    "    '40', '50', '55', '60', '65', '70', '75', '80',\n"
    "}\n\n"
    "_OCR_FIXES = {\n"
    "    'O': '0', 'I': '1', 'Z': '2',\n"
    "    'S': '5', 'B': '8', 'G': '6', 'Q': '0',\n"
    "}\n\n"
    "def validate_plate(text: str) -> tuple[bool, str]:\n"
    "    normalized = normalize(text)\n"
    "    if UZ_PLATE_RE.match(normalized) and normalized[:2] in VALID_REGIONS:\n"
    "        return True, normalized\n"
    "    corrected = _try_correct(normalized)\n"
    "    if corrected and UZ_PLATE_RE.match(corrected) and corrected[:2] in VALID_REGIONS:\n"
    "        return True, corrected\n"
    "    return False, normalized"
)

body_para("Image preprocessing pipeline (core/recognizer.py):", bold=True, indent_first=0)

code_block(
    "def _preprocess(self, crop: np.ndarray) -> np.ndarray:\n"
    "    if len(crop.shape) == 3:\n"
    "        gray = cv2.cvtColor(crop, cv2.COLOR_BGR2GRAY)\n"
    "    else:\n"
    "        gray = crop.copy()\n"
    "    scale = _TARGET_HEIGHT / gray.shape[0]\n"
    "    new_w = max(1, int(gray.shape[1] * scale))\n"
    "    resized = cv2.resize(gray, (new_w, _TARGET_HEIGHT),\n"
    "                         interpolation=cv2.INTER_CUBIC)\n"
    "    _, binary = cv2.threshold(\n"
    "        resized, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU\n"
    "    )\n"
    "    kernel = np.array([[0,-1,0],[-1,5,-1],[0,-1,0]], dtype=np.float32)\n"
    "    return cv2.filter2D(binary, -1, kernel)"
)

body_para("Database schema (db/schema.sql):", bold=True, indent_first=0)

code_block(
    "CREATE TABLE IF NOT EXISTS detections (\n"
    "    id          INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "    source      TEXT    NOT NULL,\n"
    "    plate_text  TEXT    NOT NULL,\n"
    "    raw_text    TEXT    NOT NULL,\n"
    "    is_valid    INTEGER NOT NULL DEFAULT 0,\n"
    "    confidence  REAL    NOT NULL,\n"
    "    bbox_x1     INTEGER,\n"
    "    bbox_y1     INTEGER,\n"
    "    bbox_x2     INTEGER,\n"
    "    bbox_y2     INTEGER,\n"
    "    detected_at TEXT    NOT NULL\n"
    ");\n\n"
    "CREATE TABLE IF NOT EXISTS valid_plates (\n"
    "    id          INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "    plate_text  TEXT    NOT NULL,\n"
    "    confidence  REAL    NOT NULL,\n"
    "    source      TEXT    NOT NULL,\n"
    "    detected_at TEXT    NOT NULL\n"
    ");"
)

# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────

output_path = "/home/claude/license/License_Plate_Recognition_Thesis.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
